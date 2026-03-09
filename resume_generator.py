#!/usr/bin/env python3
"""
Resume Generator Module
=======================
Takes job JSON(s) and generates tailored resumes with ATS retry loop.

Standalone:
    python resume_generator.py job.json                    # Single job
    python resume_generator.py jobs.json                   # Multiple jobs (array)
    python resume_generator.py jobs.json --top 5           # Only top 5
    python resume_generator.py job.json --no-ats-retry     # Skip retry loop

Pipeline:
    from resume_generator import generate_resume, generate_all
    path = generate_resume(job_dict, config)
    paths = generate_all(jobs_list, config)
"""

import os
import re
import copy
import json
import yaml
import logging
import argparse
from pathlib import Path
from collections import Counter

from job_schema import validate_job

logger = logging.getLogger(__name__)
BASE_DIR = Path(__file__).parent

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openai import OpenAI
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass


# ============================================================
# TITLE CLEANING
# ============================================================

def clean_job_title(raw_title: str) -> str:
    """Extract clean job title. General rules, no hardcoding."""
    title = raw_title.strip()
    # Decode HTML entities: &amp;#8211; -> –, &amp; -> &
    import html
    title = html.unescape(title)
    title = re.sub(r'&amp;?#?\w+;?', ' ', title)  # catch any remaining
    title = re.sub(r'(?i)^SR\.?\s+', 'Senior ', title)
    title = re.sub(r'(?i)^JR\.?\s+', 'Junior ', title)
    title = re.sub(r'\s*\([^)]*\)\s*', ' ', title)
    title = re.sub(r'\s*[-–—]\s*\S.*$', '', title)
    title = re.sub(r'\s*,\s+.*$', '', title)
    title = re.sub(r'(?i)\bfront[\s-]end\b', 'Frontend', title)
    title = re.sub(r'(?i)\bback[\s-]end\b', 'Backend', title)
    title = re.sub(r'(?i)\bfull[\s-]stack\b', 'Fullstack', title)
    title = re.sub(r'\s+', ' ', title).strip()
    roles = ["developer", "engineer", "architect", "designer", "analyst", "consultant",
             "administrator", "manager", "lead", "specialist", "programmer", "devops", "sre"]
    if not any(r in title.lower() for r in roles) or len(title) < 5:
        return "Full-Stack Developer"
    return title


def sanitize_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = re.sub(r'\s+', '_', name.strip())
    return name[:80]


# ============================================================
# GPT TAILORING
# ============================================================

_client = None

def _get_client():
    global _client
    if _client:
        return _client
    if not HAS_OPENAI:
        return None
    key = os.getenv("OPENAI_API_KEY", "")
    if not key or key == "sk-your-key-here":
        return None
    _client = OpenAI(api_key=key)
    return _client


TECH_TERMS = {
    # Languages
    "python", "javascript", "typescript", "java", "sql", "php", "ruby", "golang",
    "rust", "scala", "kotlin", "swift", "perl", "bash", "shell", "groovy",
    "c++", "c#", ".net", "json", "xml", "yaml", "html", "css",
    # Frontend
    "react", "react.js", "vue", "vue.js", "angular", "svelte", "next.js", "nuxt",
    "sass", "less", "tailwind", "bootstrap", "webpack", "vite", "babel",
    "redux", "mobx", "graphql", "apollo", "jquery", "ember",
    # Backend
    "node", "node.js", "express", "django", "flask", "fastapi", "spring",
    "springboot", "rails", "laravel", "nestjs", "nest.js", "sinatra",
    "rest", "restful", "api", "apis", "grpc", "websocket",
    "microservices", "microservice",
    # Databases
    "postgresql", "postgres", "mysql", "mongodb", "redis", "elasticsearch",
    "dynamodb", "cassandra", "sqlite", "oracle", "mariadb", "couchdb",
    "neo4j", "memcached",
    # Message Queues / Streaming
    "kafka", "rabbitmq", "activemq", "redis", "celery", "sidekiq",
    "mqtt", "nats", "pulsar",
    # Cloud/DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
    "jenkins", "bamboo", "gradle", "maven", "ant",
    "cicd", "github", "gitlab", "bitbucket",
    "linux", "nginx", "apache", "tomcat",
    "kubectl", "helm", "istio", "prometheus", "grafana", "datadog",
    "cloudformation", "pulumi", "vagrant",
    "eks", "ecs", "fargate", "lambda",
    # Tools/Practices
    "git", "svn", "agile", "scrum", "kanban", "jira", "confluence",
    "devops", "sre", "monitoring", "logging", "observability",
    "testing", "tdd", "bdd", "jest", "cypress", "selenium", "pytest",
    "junit", "mocha", "chai", "playwright", "puppeteer",
    # Java ecosystem
    "jdk", "jvm", "hibernate", "mybatis", "tomcat", "wildfly",
    "lombok", "jackson",
    # CMS/Web
    "wordpress", "drupal", "shopify", "woocommerce", "seo", "wcag",
    "accessibility", "responsive",
    # Data
    "spark", "hadoop", "airflow", "etl", "dbt", "snowflake",
    "bigquery", "redshift", "databricks", "pandas", "numpy",
    # Architecture
    "serverless", "containers", "cdn", "caching",
    "oauth", "jwt", "authentication", "authorization",
    "saml", "sso", "ldap",
    # Soft-technical
    "fullstack", "frontend", "backend",
    "architecture", "scalable", "scalability",
    "performance", "optimization", "security",
}

# Normalization map: raw text patterns → standardized form
# Applied before word extraction so hyphenated/spaced terms get matched
TEXT_NORMALIZATIONS = {
    "front-end": "frontend",
    "front end": "frontend",
    "back-end": "backend",
    "back end": "backend",
    "full-stack": "fullstack",
    "full stack": "fullstack",
    "ci/cd": "cicd",
    "ci / cd": "cicd",
    "spring boot": "springboot",
    "spring-boot": "springboot",
    "node.js": "node.js",
    "vue.js": "vue.js",
    "react.js": "react.js",
    "next.js": "next.js",
    "nest.js": "nest.js",
    "e2e": "testing",
    "end-to-end": "testing",
}


def _extract_top_keywords(text: str, n: int = 25, company: str = "") -> list[tuple[str, int]]:
    """Extract only technical/skill keywords from job description for GPT prompt."""
    # Clean HTML entities, URLs, hashtags
    text = re.sub(r'&amp;?#?\w+;?', ' ', text)
    text = re.sub(r'https?://\S+', ' ', text)
    text = re.sub(r'#\w+', ' ', text)
    text = re.sub(r'[<>{}]', ' ', text)
    
    # Normalize compound terms BEFORE word extraction
    text_lower = text.lower()
    for raw, normalized in TEXT_NORMALIZATIONS.items():
        text_lower = text_lower.replace(raw, normalized)
    
    # Extract words (preserve tech terms like C++, C#, Node.js)
    words = re.findall(r'[a-zA-Z][a-zA-Z\+\#\.]*[a-zA-Z\+\#]|[a-zA-Z]{3,}', text_lower)
    words = [w.rstrip('.') for w in words]
    
    counts = Counter(words)
    
    # Only return words that are in our tech terms list or have special chars (C++, C#, Node.js)
    tech_keywords = []
    for word, count in counts.most_common(100):
        if word in TECH_TERMS or any(c in word for c in ['+', '#', '.']):
            tech_keywords.append((word, count))
    
    return tech_keywords[:n]


def tailor_for_job(job: dict, config: dict, missing_keywords: list[str] = None) -> dict | None:
    """Call GPT to tailor resume for a specific job. Returns tailored data dict."""
    client = _get_client()
    if not client:
        logger.warning("No OpenAI client. Skipping tailoring.")
        return None

    base_resume = config["base_resume"]
    profile = config["profile"]
    job_desc = job.get("description", "")
    job_title = job.get("title", "")
    company = job.get("company", "")
    clean_title = clean_job_title(job_title)
    matched = job.get("matched_skills", [])

    top_kw = _extract_top_keywords(job_desc, company=company)
    kw_str = ", ".join(f"{kw}({c}x)" for kw, c in top_kw if c >= 2)

    # Load prompt templates from files
    prompts_dir = Path(__file__).parent / "prompts"
    
    system_prompt = (prompts_dir / "system.txt").read_text(encoding="utf-8").strip()
    tailor_template = (prompts_dir / "tailor.txt").read_text(encoding="utf-8")
    
    # Build missing keywords section for retries
    missing_str = ""
    if missing_keywords:
        retry_template = (prompts_dir / "retry.txt").read_text(encoding="utf-8")
        missing_str = retry_template.replace("{missing_keywords}", ", ".join(missing_keywords[:15]))

    # Build bullet strings
    bullets_city = chr(10).join(f"- {b}" for b in base_resume["bullets"].get("city_of_gatineau", []))
    bullets_pos = chr(10).join(f"- {b}" for b in base_resume["bullets"].get("precision_os", []))
    bullets_syn = chr(10).join(f"- {b}" for b in base_resume["bullets"].get("syntax", []))
    bullets_uot = chr(10).join(f"- {b}" for b in base_resume["bullets"].get("uottawa", []))

    prompt = tailor_template.format(
        clean_title=clean_title,
        company=company,
        job_desc=job_desc[:2500],
        kw_str=kw_str,
        missing_str=missing_str,
        bullets_city_of_gatineau=bullets_city,
        bullets_precision_os=bullets_pos,
        bullets_syntax=bullets_syn,
        bullets_uottawa=bullets_uot,
    )

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4, max_tokens=3000,
        )
        content = resp.choices[0].message.content.strip()
        if content.startswith("```"):
            content = content.split("\n", 1)[1] if "\n" in content else content[3:]
        if content.endswith("```"):
            content = content[:-3]
        result = json.loads(content.strip())
        logger.info(f"Tailored: {job_title} @ {company}")
        return result
    except Exception as e:
        logger.error(f"Tailoring failed: {e}")
        return None


# ============================================================
# TEMPLATE FILLING
# ============================================================

DEFAULT_SKILLS = {
    "Languages": "JavaScript, TypeScript, Python, Java, C++, SQL, PHP",
    "Frontend": "React.js, HTML/CSS, Tailwind, Bootstrap, Responsive Design",
    "Backend": "Node.js, Express, Spring Boot, REST APIs, .NET",
    "Databases": "PostgreSQL, MongoDB, SQL Server",
    "Cloud & DevOps": "Azure, Docker, CI/CD, Git, SAP Cloud Platform",
    "Tools": "GitHub Copilot, AI-assisted development, Agile/Scrum, Jira",
    "Spoken Languages": "French (Fluent, C2), English (Fluent, C2)",
}


def _get_tailored_skills(tailored_data: dict) -> dict:
    skills = tailored_data.get("skills", {})
    if not skills or not isinstance(skills, dict):
        return DEFAULT_SKILLS.copy()
    has_spoken = any("french" in v.lower() for v in skills.values())
    if not has_spoken:
        skills["Spoken Languages"] = "French (Fluent, C2), English (Fluent, C2)"
    return skills


def _normalize_key(key: str) -> str:
    return re.sub(r'[^a-z]', '', key.lower())


def _get_bullets(tailored_bullets: dict, role_key: str, base_bullets: list) -> list:
    # Exact match
    b = tailored_bullets.get(role_key)
    if b:
        cleaned = [x for x in b if x and x.strip()]
        if cleaned:
            return cleaned
    # Fuzzy match
    target = _normalize_key(role_key)
    for k, v in tailored_bullets.items():
        if _normalize_key(k) == target or target in _normalize_key(k) or _normalize_key(k) in target:
            cleaned = [x for x in v if x and x.strip()]
            if cleaned:
                return cleaned
    return base_bullets or []


def _fill_template(doc, replacements, skills_keys=None, job_header_keys=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    skills_keys = skills_keys or []
    job_header_keys = job_header_keys or []

    for para in doc.paragraphs:
        for key, value in replacements.items():
            ph = "{{" + key + "}}"
            if ph not in para.text:
                continue
            if key in skills_keys and ":" in value:
                label, content = value.split(":", 1)
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = label + ":"
                    para.runs[0].bold = True
                    nr = copy.deepcopy(para.runs[0]._element)
                    para._element.append(nr)
                    para.runs[-1].text = content
                    para.runs[-1].bold = False
                continue
            if key in job_header_keys and "|" in value:
                parts = value.split("|", 1)
                title_part = parts[0].strip()
                rest = parts[1].strip()
                company_part, dates_part = (rest.split("\t", 1) + [""])[:2]
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = title_part + " | "
                    para.runs[0].bold = True
                    para.runs[0].italic = False
                    r2 = copy.deepcopy(para.runs[0]._element)
                    para._element.append(r2)
                    para.runs[-1].text = company_part.strip()
                    para.runs[-1].bold = False
                    para.runs[-1].italic = True
                    if dates_part:
                        r3 = copy.deepcopy(para.runs[0]._element)
                        para._element.append(r3)
                        para.runs[-1].text = "\t" + dates_part.strip()
                        para.runs[-1].bold = False
                        para.runs[-1].italic = False
                continue
            for run in para.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, value)
            if ph in para.text:
                full = para.text
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = full
                    for k2, v2 in replacements.items():
                        p2 = "{{" + k2 + "}}"
                        if p2 in para.runs[0].text:
                            para.runs[0].text = para.runs[0].text.replace(p2, v2)


def _add_hyperlink(doc, display_text, url):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for para in doc.paragraphs:
        if display_text not in para.text:
            continue
        for run in para.runs:
            if display_text not in run.text:
                continue
            before, after = run.text.split(display_text, 1)
            run.text = before
            r_id = doc.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hl = OxmlElement('w:hyperlink')
            hl.set(qn('r:id'), r_id)
            nr = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
            u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
            nr.append(rPr)
            t = OxmlElement('w:t'); t.text = display_text; nr.append(t)
            hl.append(nr)
            run._element.addnext(hl)
            if after:
                ar = copy.deepcopy(run._element)
                for ct in ar.findall(qn('w:t')):
                    ct.text = after
                hl.addnext(ar)
            return


# ============================================================
# RESUME GENERATION
# ============================================================

def generate_resume(job: dict, config: dict, tailored_data: dict = None,
                    output_dir: Path = None) -> Path | None:
    """Generate a tailored resume for one job. Returns filepath."""
    if not HAS_DOCX:
        logger.error("python-docx not installed")
        return None

    profile = config["profile"]
    base_resume = config["base_resume"]
    tailored_data = tailored_data or {}
    output_dir = output_dir or Path("output/resumes")
    output_dir.mkdir(parents=True, exist_ok=True)

    template_path = config.get("tailoring", {}).get("template", "")
    tp = None
    if template_path:
        tp = Path(template_path)
        if not tp.exists():
            tp = BASE_DIR / template_path
        if not tp.exists():
            logger.warning(f"Template not found: {template_path}")
            tp = None

    bullets_data = tailored_data.get("bullets", {})
    def get_b(key):
        return _get_bullets(bullets_data, key, base_resume["bullets"].get(key, []))
    def bof(bullets, idx, base_key):
        if idx < len(bullets) and bullets[idx] and bullets[idx].strip():
            return bullets[idx]
        base = base_resume["bullets"].get(base_key, [])
        return base[idx] if idx < len(base) else ""

    summary = tailored_data.get("summary", (
        f"Full-Stack Developer with {profile['years_experience']}+ years of experience "
        "building scalable web applications, REST APIs, and cloud-based solutions."))

    skills = _get_tailored_skills(tailored_data)
    skills_items = list(skills.items())
    ct = clean_job_title(job.get("title", "Full-Stack Developer"))

    g = get_b("city_of_gatineau")
    p = get_b("precision_os")
    s = get_b("syntax")
    u = get_b("uottawa")

    skills_phs = ["SKILLS_LANGUAGES", "SKILLS_FRAMEWORKS", "SKILLS_WEB",
                  "SKILLS_CLOUD", "SKILLS_DEVOPS", "SKILLS_AI", "SKILLS_SPOKEN"]

    repl = {
        "NAME": profile["name"], "TITLE": ct,
        "CONTACT": f"{profile['email']} | LinkedIn",
        "SUMMARY": summary,
    }
    for i, ph in enumerate(skills_phs):
        if i < len(skills_items):
            label, content = skills_items[i]
            repl[ph] = f"{label}: {content}"
        else:
            repl[ph] = ""

    repl.update({
        "JOB1_HEADER": f"Full-Stack Developer | City of Gatineau\tSeptember 2023 \u2013 Present",
        "JOB1_BULLET_1": bof(g, 0, "city_of_gatineau"), "JOB1_BULLET_2": bof(g, 1, "city_of_gatineau"), "JOB1_BULLET_3": bof(g, 2, "city_of_gatineau"),
        "JOB2_HEADER": f"Software Engineer | Precision OS\tJanuary 2022 \u2013 September 2023",
        "JOB2_BULLET_1": bof(p, 0, "precision_os"), "JOB2_BULLET_2": bof(p, 1, "precision_os"), "JOB2_BULLET_3": bof(p, 2, "precision_os"),
        "JOB3_HEADER": f"Full-Stack Developer | Syntax (Consulting)\tJanuary 2021 \u2013 January 2022",
        "JOB3_BULLET_1": bof(s, 0, "syntax"), "JOB3_BULLET_2": bof(s, 1, "syntax"), "JOB3_BULLET_3": bof(s, 2, "syntax"),
        "JOB4_HEADER": f"Web Developer | University of Ottawa\tMay 2019 \u2013 January 2021",
        "JOB4_BULLET_1": bof(u, 0, "uottawa"), "JOB4_BULLET_2": bof(u, 1, "uottawa"), "JOB4_BULLET_3": bof(u, 2, "uottawa"),
        "EDUCATION": "Bachelor of Applied Science, Computer Engineering\t2016 \u2013 2020",
        "UNIVERSITY": "University of Ottawa",
    })

    if tp:
        doc = Document(str(tp))
        _fill_template(doc, repl, skills_keys=skills_phs,
                       job_header_keys=["JOB1_HEADER", "JOB2_HEADER", "JOB3_HEADER", "JOB4_HEADER"])
    else:
        doc = _build_from_scratch(repl, profile, base_resume, tailored_data, ct)

    # Hyperlink
    linkedin = profile.get("linkedin", "")
    if linkedin:
        _add_hyperlink(doc, "LinkedIn", linkedin)

    # Save in folder
    company = job.get("company", "Unknown")
    safe_co = sanitize_filename(company)
    safe_ti = sanitize_filename(ct)
    folder = output_dir / f"{safe_co}_{safe_ti}"
    folder.mkdir(exist_ok=True)

    filepath = folder / "AG_Resume.docx"
    doc.save(str(filepath))

    # Save job link
    job_url = job.get("url", "")
    if job_url:
        (folder / "job_link.txt").write_text(f"{job.get('title','')}\n{company}\n{job_url}\n", encoding="utf-8")

    logger.info(f"Generated: {folder.name}/AG_Resume.docx")
    return filepath


def _build_from_scratch(repl, profile, base_resume, tailored_data, ct):
    """Build resume from scratch when no template."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(profile["name"].upper())
    r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ct)
    r.font.size = Pt(12); r.font.name = "Calibri"

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{profile['email']} | LinkedIn")
    r.font.size = Pt(9); r.font.name = "Calibri"

    # Summary
    _add_section(doc, "PROFESSIONAL SUMMARY")
    doc.add_paragraph(repl.get("SUMMARY", ""))

    # Skills
    _add_section(doc, "TECHNICAL SKILLS")
    skills = _get_tailored_skills(tailored_data)
    for label, vals in skills.items():
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: "); r.bold = True; r.font.size = Pt(10)
        p.add_run(vals).font.size = Pt(10)

    # Experience
    _add_section(doc, "WORK EXPERIENCE")
    roles = [
        ("Full-Stack Developer", "City of Gatineau", "September 2023 – Present", "city_of_gatineau"),
        ("Software Engineer", "Precision OS", "January 2022 – September 2023", "precision_os"),
        ("Full-Stack Developer", "Syntax (Consulting)", "January 2021 – January 2022", "syntax"),
        ("Web Developer", "University of Ottawa", "May 2019 – January 2021", "uottawa"),
    ]
    bullets_data = tailored_data.get("bullets", {})
    for title, co, dates, key in roles:
        p = doc.add_paragraph()
        r = p.add_run(f"{title} | {co}"); r.bold = True; r.font.size = Pt(10.5)
        p.add_run(f"    {dates}").font.size = Pt(10)
        for b in _get_bullets(bullets_data, key, base_resume["bullets"].get(key, [])):
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(b).font.size = Pt(10)

    # Education
    _add_section(doc, "EDUCATION")
    p = doc.add_paragraph()
    r = p.add_run("Bachelor of Applied Science, Computer Engineering"); r.bold = True
    p.add_run("    2016 – 2020")
    doc.add_paragraph("University of Ottawa")

    return doc


def _add_section(doc, text):
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    p.space_before = Pt(12); p.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '000000'})
    pBdr.append(bottom); pPr.append(pBdr)


# ============================================================
# ATS RETRY LOOP
# ============================================================

def generate_with_ats_retry(job: dict, config: dict, output_dir: Path = None) -> Path | None:
    """Generate resume, check ATS score, retry if below threshold."""
    from ats_checker import validate_resume_file, format_report

    max_retries = config.get("resume_generator", {}).get("max_retries", 3)
    min_keyword_match = config.get("resume_generator", {}).get("min_keyword_pct", 45)
    job_desc = job.get("description", "")

    tailored = tailor_for_job(job, config)
    missing_kw = []

    for attempt in range(1, max_retries + 1):
        if attempt > 1:
            logger.info(f"  Retry {attempt}/{max_retries} - adding missing keywords: {', '.join(missing_kw[:10])}")
            tailored = tailor_for_job(job, config, missing_keywords=missing_kw)

        filepath = generate_resume(job, config, tailored, output_dir)
        if not filepath:
            return None

        if not job_desc:
            break

        result = validate_resume_file(str(filepath), job_desc, company=job.get("company", ""))
        km = result.get("keyword_match", {})
        match_pct = km.get("match_pct", 0)
        missing_kw = km.get("missing_keywords", [])

        # Save ATS report
        ats_path = filepath.parent / "ats_report.txt"
        ats_path.write_text(format_report(result), encoding="utf-8")

        logger.info(f"  ATS: {match_pct}% keyword match (target: {min_keyword_match}%)")

        if match_pct >= min_keyword_match or not missing_kw:
            break

    return filepath


# ============================================================
# BATCH GENERATION
# ============================================================

def generate_all(jobs: list[dict], config: dict, output_dir: Path = None) -> list[Path]:
    """Generate resumes for multiple jobs with ATS retry."""
    output_dir = output_dir or Path("output/resumes")
    generated = []
    top_n = config.get("tailoring", {}).get("top_n", 15)
    min_score = config.get("tailoring", {}).get("min_score", 0.3)

    eligible = [j for j in jobs[:top_n] if j.get("score", 1.0) >= min_score]
    logger.info(f"Generating resumes for {len(eligible)} jobs")

    for job in eligible:
        path = generate_with_ats_retry(job, config, output_dir)
        if path:
            generated.append(path)

    logger.info(f"Generated {len(generated)} resumes")
    return generated


# ============================================================
# CLI
# ============================================================

def main():
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s", datefmt="%H:%M:%S")
    parser = argparse.ArgumentParser(description="Generate tailored resumes from job JSON")
    parser.add_argument("jobs", help="Path to job.json (single) or jobs.json (array)")
    parser.add_argument("-c", "--config", default=str(BASE_DIR / "config.yaml"))
    parser.add_argument("-o", "--output", default="output/resumes")
    parser.add_argument("--top", type=int, default=0, help="Only process top N jobs")
    parser.add_argument("--no-ats-retry", action="store_true")
    args = parser.parse_args()

    with open(args.config) as f:
        config = yaml.safe_load(f)

    with open(args.jobs) as f:
        data = json.load(f)

    jobs = data if isinstance(data, list) else [data]
    if args.top > 0:
        jobs = jobs[:args.top]

    out = Path(args.output)
    if args.no_ats_retry:
        for job in jobs:
            tailored = tailor_for_job(job, config)
            generate_resume(job, config, tailored, out)
    else:
        generate_all(jobs, config, out)


if __name__ == "__main__":
    main()