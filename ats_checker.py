"""
ATS Resume Validator
====================
Checks a .docx resume for ATS compatibility issues.
Can be used standalone or integrated into the pipeline.

Standalone usage:
    python ats_checker.py path/to/resume.docx
    python ats_checker.py path/to/resume.docx --job-description "paste job desc here"
    python ats_checker.py path/to/resume.docx --job-file job_description.txt

Pipeline usage:
    from ats_checker import validate_resume, validate_resume_file
    result = validate_resume_file("resume.docx")
    result = validate_resume_file("resume.docx", job_description="...")
"""

import sys
import re
import logging
from pathlib import Path

# Import shared tech terms and normalizations from resume_generator
try:
    from resume_generator import TECH_TERMS, TEXT_NORMALIZATIONS
except ImportError:
    TECH_TERMS = set()
    TEXT_NORMALIZATIONS = {}
from collections import Counter

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================
# ATS CHECKS
# ============================================================

def check_file_format(filepath: str) -> dict:
    """Check if file is a valid .docx."""
    path = Path(filepath)
    issues = []
    
    if not path.exists():
        return {"pass": False, "issues": [f"File not found: {filepath}"], "severity": "critical"}
    
    if path.suffix.lower() not in [".docx"]:
        issues.append(f"File is {path.suffix}, not .docx. Many ATS systems reject .doc, .pdf generated from Word, etc.")
        return {"pass": False, "issues": issues, "severity": "critical"}
    
    try:
        Document(str(path))
    except Exception as e:
        issues.append(f"File is corrupted or not a valid .docx: {e}")
        return {"pass": False, "issues": issues, "severity": "critical"}
    
    return {"pass": True, "issues": [], "severity": "ok"}


def check_text_extractable(doc: Document) -> dict:
    """Check if all text can be extracted (ATS reads plain text)."""
    issues = []
    full_text = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    if not full_text:
        issues.append("No extractable text found. ATS will see a blank resume.")
        return {"pass": False, "issues": issues, "severity": "critical", "text": ""}
    
    combined = "\n".join(full_text)
    
    if len(combined) < 200:
        issues.append(f"Very little text extracted ({len(combined)} chars). Resume may be mostly in images or text boxes.")
    
    return {"pass": len(issues) == 0, "issues": issues, "severity": "warning" if issues else "ok", "text": combined}


def check_fonts(doc: Document) -> dict:
    """Check for ATS-safe fonts."""
    issues = []
    fonts_used = set()
    
    safe_fonts = {
        "calibri", "arial", "helvetica", "times new roman", "georgia",
        "verdana", "tahoma", "trebuchet ms", "cambria", "garamond",
        "book antiqua", "palatino linotype", "century gothic",
    }
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
    
    unsafe_fonts = []
    for f in fonts_used:
        if f.lower() not in safe_fonts:
            unsafe_fonts.append(f)
    
    if unsafe_fonts:
        issues.append(f"Non-standard fonts detected: {', '.join(unsafe_fonts)}. Some ATS may not render these correctly.")
    
    return {"pass": len(issues) == 0, "issues": issues, "severity": "warning" if issues else "ok", "fonts": list(fonts_used)}


def check_tables(doc: Document) -> dict:
    """Check for tables used as layout (ATS often can't parse these)."""
    issues = []
    
    table_count = len(doc.tables)
    
    if table_count > 0:
        # Check if tables are used for layout vs data
        for i, table in enumerate(doc.tables):
            rows = len(table.rows)
            cols = len(table.columns)
            
            if cols >= 2 and rows >= 3:
                issues.append(
                    f"Table {i+1} ({rows}x{cols}) detected. Multi-column table layouts often break ATS parsing. "
                    "ATS reads left-to-right, row-by-row, which can scramble your content."
                )
            elif cols >= 2:
                issues.append(
                    f"Table {i+1} ({rows}x{cols}) detected. Two-column layouts can cause ATS to merge or scramble text."
                )
    
    return {"pass": len(issues) == 0, "issues": issues, "severity": "warning" if issues else "ok", "table_count": table_count}


def check_images(doc: Document) -> dict:
    """Check for images (ATS can't read text in images)."""
    issues = []
    image_count = 0
    
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    
    if image_count > 0:
        issues.append(
            f"{image_count} image(s) found. ATS cannot read text in images. "
            "If your name, contact info, or skills are in images, ATS won't see them."
        )
    
    return {"pass": image_count == 0, "issues": issues, "severity": "warning" if issues else "ok", "image_count": image_count}


def check_headers_footers(doc: Document) -> dict:
    """Check if important content is in headers/footers (some ATS skip these)."""
    issues = []
    header_text = []
    footer_text = []
    
    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                if para.text.strip():
                    header_text.append(para.text.strip())
        if section.footer:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    footer_text.append(para.text.strip())
    
    if header_text:
        issues.append(
            f"Text in header: '{' | '.join(header_text)[:100]}'. "
            "Some ATS systems skip headers. Make sure critical info (name, contact) is also in the body."
        )
    
    if footer_text:
        issues.append(
            f"Text in footer: '{' | '.join(footer_text)[:100]}'. "
            "Some ATS systems skip footers."
        )
    
    return {"pass": len(issues) == 0, "issues": issues, "severity": "info" if issues else "ok"}


def check_text_boxes(doc: Document) -> dict:
    """Check for text boxes (ATS often ignores these completely)."""
    issues = []
    textbox_count = 0
    
    # Text boxes appear as w:txbxContent in the XML
    for element in doc.element.iter():
        if element.tag.endswith('}txbxContent'):
            textbox_count += 1
    
    if textbox_count > 0:
        issues.append(
            f"{textbox_count} text box(es) found. Most ATS systems completely ignore text boxes. "
            "Any content inside them will be invisible to ATS."
        )
    
    return {"pass": textbox_count == 0, "issues": issues, "severity": "critical" if textbox_count else "ok", "textbox_count": textbox_count}


def check_sections(text: str) -> dict:
    """Check if resume has expected ATS-parseable sections."""
    issues = []
    
    section_patterns = {
        "contact": r"(email|phone|@|linkedin)",
        "experience": r"(experience|work history|employment|professional)",
        "education": r"(education|university|bachelor|master|degree|b\.?s\.?|m\.?s\.?)",
        "skills": r"(skills|technologies|technical|proficiencies|competencies)",
    }
    
    text_lower = text.lower()
    found_sections = {}
    
    for section, pattern in section_patterns.items():
        found_sections[section] = bool(re.search(pattern, text_lower))
    
    missing = [s for s, found in found_sections.items() if not found]
    
    if missing:
        issues.append(
            f"Missing standard sections: {', '.join(missing)}. "
            "ATS systems look for these section headers to categorize your content."
        )
    
    return {"pass": len(issues) == 0, "issues": issues, "severity": "warning" if issues else "ok", "sections_found": found_sections}


def check_special_characters(text: str) -> dict:
    """Check for characters that might break ATS parsing."""
    issues = []
    
    # Check for common problematic characters
    problems = []
    if '\u2022' in text:
        pass  # Standard bullet, usually fine
    
    fancy_chars = {
        '\u2013': 'en-dash (–)',
        '\u2014': 'em-dash (—)',
        '\u2018': 'smart single quote left',
        '\u2019': 'smart single quote right',
        '\u201c': 'smart double quote left',
        '\u201d': 'smart double quote right',
        '\u2026': 'ellipsis (…)',
    }
    
    found_fancy = []
    for char, name in fancy_chars.items():
        if char in text:
            found_fancy.append(name)
    
    if found_fancy:
        issues.append(
            f"Smart/fancy characters found: {', '.join(found_fancy[:5])}. "
            "Most modern ATS handle these fine, but older systems may show garbled text."
        )
    
    return {"pass": True, "issues": issues, "severity": "info" if issues else "ok"}


def check_file_size(filepath: str) -> dict:
    """Check file size (some ATS reject large files)."""
    issues = []
    size_mb = Path(filepath).stat().st_size / (1024 * 1024)
    
    if size_mb > 5:
        issues.append(f"File is {size_mb:.1f}MB. Many ATS systems reject files over 5MB.")
    elif size_mb > 2:
        issues.append(f"File is {size_mb:.1f}MB. Consider reducing size (remove images, compress).")
    
    return {"pass": size_mb <= 5, "issues": issues, "severity": "warning" if issues else "ok", "size_mb": round(size_mb, 2)}


def check_keyword_match(text: str, job_description: str, company: str = "") -> dict:
    """Check keyword overlap between resume and job description."""
    if not job_description:
        return {"pass": True, "issues": [], "severity": "skipped", "match_pct": None}
    
    stop_words = {
        # Common English
        "the", "and", "for", "are", "but", "not", "you", "all", "can", "her", "was",
        "one", "our", "out", "has", "have", "been", "from", "they", "with", "this",
        "that", "will", "your", "their", "about", "would", "there", "these", "other",
        "into", "more", "some", "such", "than", "them", "then", "what", "when", "which",
        "who", "how", "each", "she", "two", "way", "its", "may", "also", "must",
        "per", "any", "own", "both", "being", "after", "before", "here", "where",
        "over", "under", "most", "very", "just", "only", "did", "does", "done",
        # Job posting filler
        "work", "working", "ability", "strong", "experience", "team", "role", "looking",
        "join", "company", "position", "candidate", "ideal", "required", "preferred",
        "including", "using", "etc", "well", "good", "great", "should", "could",
        "based", "help", "make", "like", "need", "new", "high", "best", "ensure",
        "across", "within", "part", "take", "year", "years", "day", "time",
        "apply", "equal", "opportunity", "employer", "status", "employment",
        "benefits", "salary", "compensation", "base", "annually", "range",
        # Legal boilerplate
        "accommodation", "accommodations", "applicant", "applicants", "applicable",
        "discrimination", "disability", "race", "color", "religion", "gender",
        "national", "origin", "sexual", "orientation", "veteran", "protected",
        "law", "federal", "state", "local", "comply", "accordance",
        "conditions", "terms", "agree", "right", "rights", "notice",
        # Generic verbs/adjectives
        "deliver", "drive", "driven", "effective", "enable", "build", "create",
        "develop", "developing", "developed", "support", "manage", "lead",
        "provide", "maintain", "design", "designed", "implement", "improve",
        "advanced", "direct", "clear", "better", "demand", "description",
        "duties", "responsible", "responsibilities", "building", "collaborate",
        # Non-technical noise
        "amp", "don", "business", "customer", "customers", "culture",
        "diverse", "diversity", "inclusion", "belonging", "ambitious",
        "additional", "certain", "around", "alongside", "among",
        "communication", "candidates", "bring", "coverage",
        # Corporate filler
        "colleagues", "career", "countries", "committed", "employees", "employee",
        "contribute", "continuous", "define", "developers", "engineers",
        "end", "edge", "form", "free", "get", "group", "hours", "insurance",
        "activities", "commitment", "communities", "education", "empowers",
        "enables", "capacity", "consideration", "cost", "access", "care",
        # Common German words
        "der", "die", "das", "ein", "eine", "einer", "und", "ist", "mit",
        "auf", "den", "dem", "von", "wir", "uns", "fur", "für", "sie",
        "sich", "oder", "auch", "als", "bei", "wird", "nach", "nicht",
        "wie", "hat", "dein", "deine", "unser", "unsere", "zum", "zur",
        "brauchen", "suchen", "bieten", "dich", "dir",
        # Common French words
        "les", "des", "une", "est", "dans", "pour", "avec", "sur", "qui",
        "par", "pas", "sont", "nous", "vous", "mais", "tout", "cette",
        "aux", "ses", "nos", "vos", "plus", "entre",
        # Misc
        "e.g", "i.e", "etc", "via",
    }
    
    # Dynamically exclude company name words
    if company:
        for word in re.findall(r'[a-zA-Z]{3,}', company.lower()):
            stop_words.add(word)
    
    def clean_text(t):
        t = re.sub(r'&amp;?#?\w+;?', ' ', t)
        t = re.sub(r'https?://\S+', ' ', t)
        t = re.sub(r'#\w+', ' ', t)
        t = re.sub(r'[<>{}]', ' ', t)
        return t
    
    def extract_keywords(t):
        t = clean_text(t)
        # Normalize compound terms before extraction
        t_lower = t.lower()
        for raw, normalized in TEXT_NORMALIZATIONS.items():
            t_lower = t_lower.replace(raw, normalized)
        words = re.findall(r'[a-zA-Z][a-zA-Z\+\#\.]*[a-zA-Z\+\#]|[a-zA-Z]{3,}', t_lower)
        words = [w.rstrip('.') for w in words]
        return [w for w in words if w not in stop_words and len(w) >= 3]
    
    resume_words = extract_keywords(text)
    job_words = extract_keywords(job_description)
    
    if not job_words:
        return {"pass": True, "issues": [], "severity": "skipped", "match_pct": None}
    
    job_keywords = Counter(job_words)
    resume_set = set(resume_words)
    
    # Technical/skill keywords — these are what ATS actually cares about
    # Using shared TECH_TERMS from resume_generator
    
    # Find technical keywords that appear in job description
    tech_in_job = set()
    for word in job_keywords:
        if word in TECH_TERMS or any(c in word for c in ['+', '#', '.']):
            tech_in_job.add(word)
    
    # Also catch multi-word tech like "node.js" that got split
    job_text_lower = " ".join(job_words)
    for term in ["node.js", "vue.js", "react.js", "next.js", "nest.js", "spring boot",
                  "sql server", "ci/cd", "full-stack", "front-end", "back-end"]:
        if term in job_text_lower:
            tech_in_job.add(term)
    
    # If very few tech terms found, add high-frequency non-stop words as backup
    if len(tech_in_job) < 5:
        for word, count in job_keywords.most_common(15):
            if count >= 2 and len(word) >= 4:
                tech_in_job.add(word)
    
    important_job_keywords = tech_in_job
    
    if not important_job_keywords:
        important_job_keywords = set(list(job_keywords.keys())[:20])
    
    matched = important_job_keywords & resume_set
    missing = important_job_keywords - resume_set
    
    match_pct = len(matched) / len(important_job_keywords) * 100 if important_job_keywords else 0
    
    issues = []
    
    if match_pct < 40:
        issues.append(
            f"LOW keyword match: {match_pct:.0f}%. "
            f"Missing key terms: {', '.join(sorted(missing)[:15])}. "
            "Most ATS filter at 40-60% keyword match."
        )
    elif match_pct < 60:
        issues.append(
            f"Moderate keyword match: {match_pct:.0f}%. "
            f"Consider adding: {', '.join(sorted(missing)[:10])}."
        )
    
    return {
        "pass": match_pct >= 40,
        "issues": issues,
        "severity": "critical" if match_pct < 40 else ("warning" if match_pct < 60 else "ok"),
        "match_pct": round(match_pct, 1),
        "matched_keywords": sorted(matched),
        "missing_keywords": sorted(missing)[:20],
    }


# ============================================================
# MAIN VALIDATOR
# ============================================================

def validate_resume(doc: Document, filepath: str, job_description: str = "", company: str = "") -> dict:
    """Run all ATS checks on a loaded document. Returns full report."""
    results = {}
    
    results["file_format"] = check_file_format(filepath)
    results["file_size"] = check_file_size(filepath)
    results["text_extract"] = check_text_extractable(doc)
    results["fonts"] = check_fonts(doc)
    results["tables"] = check_tables(doc)
    results["images"] = check_images(doc)
    results["headers_footers"] = check_headers_footers(doc)
    results["text_boxes"] = check_text_boxes(doc)
    results["special_chars"] = check_special_characters(results["text_extract"].get("text", ""))
    results["sections"] = check_sections(results["text_extract"].get("text", ""))
    
    if job_description:
        results["keyword_match"] = check_keyword_match(results["text_extract"].get("text", ""), job_description, company=company)
    
    # Overall score
    total_checks = 0
    passed_checks = 0
    critical_fails = []
    warnings = []
    
    for name, result in results.items():
        if result.get("severity") == "skipped":
            continue
        total_checks += 1
        if result["pass"]:
            passed_checks += 1
        elif result.get("severity") == "critical":
            critical_fails.extend(result["issues"])
        else:
            warnings.extend(result["issues"])
    
    score = (passed_checks / total_checks * 100) if total_checks > 0 else 0
    
    results["overall"] = {
        "score": round(score),
        "passed": passed_checks,
        "total": total_checks,
        "critical_fails": critical_fails,
        "warnings": warnings,
        "ats_ready": len(critical_fails) == 0 and score >= 70,
    }
    
    return results


def validate_resume_file(filepath: str, job_description: str = "", company: str = "") -> dict:
    """Validate a .docx file from path. Convenience wrapper."""
    if not HAS_DOCX:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    
    path = Path(filepath)
    if not path.exists():
        return {"error": f"File not found: {filepath}"}
    
    doc = Document(str(path))
    return validate_resume(doc, filepath, job_description, company=company)


def format_report(results: dict) -> str:
    """Format validation results as readable text."""
    lines = []
    overall = results.get("overall", {})
    
    score = overall.get("score", 0)
    emoji = "✅" if overall.get("ats_ready") else ("⚠️" if score >= 50 else "❌")
    
    lines.append(f"\n{'='*60}")
    lines.append(f"  ATS COMPATIBILITY REPORT")
    lines.append(f"{'='*60}")
    lines.append(f"\n  {emoji} Overall Score: {score}% ({overall.get('passed', 0)}/{overall.get('total', 0)} checks passed)")
    lines.append(f"  ATS Ready: {'YES' if overall.get('ats_ready') else 'NO'}")
    
    if overall.get("critical_fails"):
        lines.append(f"\n  🔴 CRITICAL ISSUES:")
        for issue in overall["critical_fails"]:
            lines.append(f"     • {issue}")
    
    if overall.get("warnings"):
        lines.append(f"\n  🟡 WARNINGS:")
        for issue in overall["warnings"]:
            lines.append(f"     • {issue}")
    
    # Keyword match details
    km = results.get("keyword_match", {})
    if km.get("match_pct") is not None:
        lines.append(f"\n  📊 Keyword Match: {km['match_pct']}%")
        if km.get("matched_keywords"):
            lines.append(f"     Matched: {', '.join(km['matched_keywords'][:15])}")
        if km.get("missing_keywords"):
            lines.append(f"     Missing: {', '.join(km['missing_keywords'][:15])}")
    
    if not overall.get("critical_fails") and not overall.get("warnings"):
        lines.append(f"\n  ✅ No issues found. Resume looks ATS-ready.")
    
    lines.append(f"\n{'='*60}\n")
    
    return "\n".join(lines)


# ============================================================
# CLI
# ============================================================

def main():
    """CLI entry point for standalone usage."""
    import argparse
    
    parser = argparse.ArgumentParser(description="ATS Resume Validator")
    parser.add_argument("resume", help="Path to .docx resume file")
    parser.add_argument("--job-description", "-j", help="Job description text to match against", default="")
    parser.add_argument("--job-file", "-f", help="Path to .txt file containing job description", default="")
    parser.add_argument("--job-json", help="Path to job.json (uses description field)", default="")
    parser.add_argument("--json", action="store_true", help="Output results as JSON")
    
    args = parser.parse_args()
    
    job_desc = args.job_description
    if args.job_json:
        import json as json_mod
        with open(args.job_json, "r", encoding="utf-8") as jf:
            data = json_mod.load(jf)
        if isinstance(data, list):
            data = data[0]
        job_desc = data.get("description", "")
    elif args.job_file:
        job_desc = Path(args.job_file).read_text(encoding="utf-8")
    
    results = validate_resume_file(args.resume, job_desc)
    
    if "error" in results:
        print(f"Error: {results['error']}")
        sys.exit(1)
    
    if args.json:
        import json as json_mod
        if "text_extract" in results:
            results["text_extract"].pop("text", None)
        print(json_mod.dumps(results, indent=2))
    else:
        print(format_report(results))


if __name__ == "__main__":
    main()