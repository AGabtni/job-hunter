"""Generate tailored .docx resumes for top job matches."""

import re
import copy
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. Run: pip install python-docx")


def sanitize_filename(name: str) -> str:
    """Make a string safe for filenames."""
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = re.sub(r'\s+', '_', name.strip())
    return name[:80]


def _normalize_key(key: str) -> str:
    """Normalize a role key for fuzzy matching."""
    return re.sub(r'[^a-z]', '', key.lower())


def _get_bullets(tailored_bullets: dict, role_key: str, base_bullets: list) -> list:
    """
    Get bullets for a role, handling GPT returning different key formats.
    Falls back to base resume bullets if tailored bullets are empty or missing.
    """
    # Try exact match first
    bullets = tailored_bullets.get(role_key)
    if bullets and len(bullets) > 0:
        bullets = [b for b in bullets if b and b.strip()]
        if bullets:
            return bullets

    # Try fuzzy match - normalize both sides
    normalized_target = _normalize_key(role_key)
    for key, val in tailored_bullets.items():
        if _normalize_key(key) == normalized_target:
            if val and len(val) > 0:
                cleaned = [b for b in val if b and b.strip()]
                if cleaned:
                    return cleaned

    # Try partial match (e.g. "gatineau" in "city_of_gatineau")
    for key, val in tailored_bullets.items():
        key_norm = _normalize_key(key)
        if key_norm in normalized_target or normalized_target in key_norm:
            if val and len(val) > 0:
                cleaned = [b for b in val if b and b.strip()]
                if cleaned:
                    return cleaned

    # Fallback to base resume bullets
    return base_bullets if base_bullets else []


# Default skills when GPT doesn't return tailored skills
DEFAULT_SKILLS = {
    "Languages": "JavaScript, TypeScript, Python, Java, C++, SQL, PHP",
    "Frontend": "React.js, HTML/CSS, Tailwind, Bootstrap, Responsive Design",
    "Backend": "Node.js, Express, Spring Boot, REST APIs, .NET",
    "Databases": "PostgreSQL, MongoDB, SQL Server",
    "Cloud & DevOps": "Azure, Docker, CI/CD, Git, SAP Cloud Platform",
    "Tools": "GitHub Copilot, AI-assisted development, Agile/Scrum, Jira",
    "Spoken Languages": "French (Fluent, C2), English (Fluent, C2)",
}


def _get_tailored_skills(tailored_data: dict) -> dict:
    """Get skills from GPT response, falling back to defaults."""
    skills = tailored_data.get("skills", {})
    if not skills or not isinstance(skills, dict):
        return DEFAULT_SKILLS.copy()
    
    # Ensure spoken languages is always present
    has_spoken = False
    for key in skills:
        if "spoken" in key.lower() or "language" in key.lower() and "programming" not in key.lower():
            if "french" in skills[key].lower():
                has_spoken = True
                break
    
    if not has_spoken:
        skills["Spoken Languages"] = "French (Fluent, C2), English (Fluent, C2)"
    
    return skills


def _clean_job_title(raw_title: str) -> str:
    """Extract clean job title from messy scraped title.
    
    Rules:
    1. Expand abbreviations (SR. -> Senior, JR. -> Junior)
    2. Strip parentheses (tech stacks, contract types, locations)
    3. Strip everything after " - " (company names, locations, tags)
    4. Strip everything after ", " (team names, qualifiers)
    5. Normalize compound words (Front End -> Frontend)
    6. Fallback to "Full-Stack Developer" if no role word found
    """
    title = raw_title.strip()
    
    # Expand abbreviations
    title = re.sub(r'(?i)^SR\.?\s+', 'Senior ', title)
    title = re.sub(r'(?i)^JR\.?\s+', 'Junior ', title)
    
    # Remove everything in parentheses
    title = re.sub(r'\s*\([^)]*\)\s*', ' ', title)
    
    # Remove everything after dash separator (with or without space)
    title = re.sub(r'\s*[-–]\s*\S.*$', '', title)
    
    # Remove trailing comma + whatever follows
    title = re.sub(r'\s*,\s+.*$', '', title)
    
    # Normalize compound words
    title = re.sub(r'(?i)\bfront[\s-]end\b', 'Frontend', title)
    title = re.sub(r'(?i)\bback[\s-]end\b', 'Backend', title)
    title = re.sub(r'(?i)\bfull[\s-]stack\b', 'Fullstack', title)
    
    # Clean whitespace
    title = re.sub(r'\s+', ' ', title).strip()
    
    # Validate: must contain a recognizable role word
    role_words = [
        "developer", "engineer", "architect", "designer", "analyst",
        "consultant", "administrator", "manager", "lead", "specialist",
        "programmer", "devops", "sre",
    ]
    if not any(r in title.lower() for r in role_words) or len(title) < 5:
        return "Full-Stack Developer"
    
    return title


def generate_resume_docx(job: dict, tailored_data: dict, config: dict, output_dir: Path) -> Path | None:
    """Generate a tailored .docx resume for a specific job.
    Uses template if configured, otherwise builds from scratch."""
    if not HAS_DOCX:
        return None
    
    template_path = config.get("tailoring", {}).get("template", "")
    
    if template_path:
        # Try absolute path first, then relative to project dir
        tp = Path(template_path)
        if not tp.exists():
            tp = Path(__file__).parent / template_path
        if not tp.exists():
            logger.warning(f"Template not found at '{template_path}' or '{tp}'. Falling back to built-in format.")
            return _generate_from_scratch(job, tailored_data, config, output_dir)
        logger.debug(f"Using template: {tp}")
        return _generate_from_template(job, tailored_data, config, output_dir, str(tp))
    else:
        return _generate_from_scratch(job, tailored_data, config, output_dir)


def _fill_template_placeholders(doc: Document, replacements: dict, skills_keys: list = None, job_header_keys: list = None):
    """Replace {{PLACEHOLDER}} tags in a Word document.
    
    Special handling:
    - skills_keys: for these keys, text before ':' is bold, after ':' is not bold
    - job_header_keys: for these keys, text after '|' (company) is italic
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    skills_keys = skills_keys or []
    job_header_keys = job_header_keys or []
    
    for para in doc.paragraphs:
        for key, value in replacements.items():
            placeholder = "{{" + key + "}}"
            if placeholder not in para.text:
                continue
            
            # Special: skills lines -> "Label: content" with bold label
            if key in skills_keys and ":" in value:
                label, content = value.split(":", 1)
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = label + ":"
                    para.runs[0].bold = True
                    new_run = copy.deepcopy(para.runs[0]._element)
                    para._element.append(new_run)
                    last_run = para.runs[-1]
                    last_run.text = content
                    last_run.bold = False
                continue
            
            # Special: job headers -> "Title | Company\tDates" with italic company
            if key in job_header_keys and "|" in value:
                parts = value.split("|", 1)
                title_part = parts[0].strip()
                rest = parts[1].strip()
                
                if "\t" in rest:
                    company_part, dates_part = rest.split("\t", 1)
                else:
                    company_part = rest
                    dates_part = ""
                
                for run in para.runs:
                    run.text = ""
                
                if para.runs:
                    para.runs[0].text = title_part + " | "
                    para.runs[0].bold = True
                    para.runs[0].italic = False
                    
                    r2 = copy.deepcopy(para.runs[0]._element)
                    para._element.append(r2)
                    run2 = para.runs[-1]
                    run2.text = company_part.strip()
                    run2.bold = False
                    run2.italic = True
                    
                    if dates_part:
                        r3 = copy.deepcopy(para.runs[0]._element)
                        para._element.append(r3)
                        run3 = para.runs[-1]
                        run3.text = "\t" + dates_part.strip()
                        run3.bold = False
                        run3.italic = False
                continue
            
            # Default: simple replacement
            for run in para.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
            
            # Handle placeholder spanning multiple runs
            if placeholder in para.text:
                full_text = para.text
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = full_text
                    for key2, value2 in replacements.items():
                        ph2 = "{{" + key2 + "}}"
                        if ph2 in para.runs[0].text:
                            para.runs[0].text = para.runs[0].text.replace(ph2, value2)


def _add_hyperlink_to_contact(doc: Document, display_text: str, url: str):
    """Find 'LinkedIn' in the contact paragraph and replace it with a clickable hyperlink."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    for para in doc.paragraphs:
        if display_text not in para.text:
            continue
        
        for run in para.runs:
            if display_text not in run.text:
                continue
            
            before, after = run.text.split(display_text, 1)
            run.text = before
            
            part = doc.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0563C1')
            rPr.append(color)
            
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)
            
            if run.font.name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), run.font.name)
                rFonts.set(qn('w:hAnsi'), run.font.name)
                rPr.append(rFonts)
            
            if run.font.size:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(run.font.size.pt * 2)))
                rPr.append(sz)
            
            new_run.append(rPr)
            
            t = OxmlElement('w:t')
            t.text = display_text
            new_run.append(t)
            
            hyperlink.append(new_run)
            run._element.addnext(hyperlink)
            
            if after:
                after_run = copy.deepcopy(run._element)
                for child_t in after_run.findall(qn('w:t')):
                    child_t.text = after
                hyperlink.addnext(after_run)
            
            return


def _generate_from_template(job: dict, tailored_data: dict, config: dict, output_dir: Path, template_path: str) -> Path | None:
    """Fill a Word template with tailored resume data."""
    profile = config["profile"]
    base_resume = config["base_resume"]
    tailored_bullets = tailored_data.get("bullets", {})
    
    doc = Document(template_path)
    
    def get_bullets(key):
        return _get_bullets(tailored_bullets, key, base_resume["bullets"].get(key, []))
    
    def bullet_or_fallback(bullets, idx, base_key):
        if idx < len(bullets) and bullets[idx] and bullets[idx].strip():
            return bullets[idx]
        base = base_resume["bullets"].get(base_key, [])
        if idx < len(base):
            return base[idx]
        return ""
    
    summary = tailored_data.get("summary", (
        f"Full-Stack Developer with {profile['years_experience']}+ years of experience "
        "building scalable web applications, REST APIs, and cloud-based solutions. "
        "Bilingual (French C2 / English C2)."
    ))
    
    # Get tailored skills from GPT (dynamic per job)
    skills = _get_tailored_skills(tailored_data)
    skills_items = list(skills.items())
    
    gatineau = get_bullets("city_of_gatineau")
    precision = get_bullets("precision_os")
    syntax = get_bullets("syntax")
    uottawa = get_bullets("uottawa")
    
    logger.debug(f"GPT returned bullet keys: {list(tailored_bullets.keys())}")
    logger.debug(f"GPT returned skills keys: {list(skills.keys())}")
    
    # Map skills to template slots (7 slots available)
    # If GPT returns fewer/more, pad or trim
    skills_placeholders = [
        "SKILLS_LANGUAGES", "SKILLS_FRAMEWORKS", "SKILLS_WEB",
        "SKILLS_CLOUD", "SKILLS_DEVOPS", "SKILLS_AI", "SKILLS_SPOKEN",
    ]
    
    # Mirror the job title in the subtitle for ATS matching
    job_title = job.get("title", "Full-Stack Developer")
    clean_title = _clean_job_title(job_title)
    # Clean up whitespace
    clean_title = re.sub(r'\s+', ' ', clean_title).strip()
    if not clean_title or len(clean_title) < 5:
        clean_title = "Full-Stack Developer"
    
    replacements = {
        "NAME": profile["name"],
        "TITLE": clean_title,
        "CONTACT": f"{profile['email']} | LinkedIn",
        "SUMMARY": summary,
    }
    
    # Fill skills slots dynamically
    for i, placeholder in enumerate(skills_placeholders):
        if i < len(skills_items):
            label, content = skills_items[i]
            replacements[placeholder] = f"{label}: {content}"
        else:
            replacements[placeholder] = ""
    
    # Job bullets
    replacements.update({
        "JOB1_HEADER": "Full-Stack Developer | City of Gatineau\tSeptember 2023 \u2013 Present",
        "JOB1_BULLET_1": bullet_or_fallback(gatineau, 0, "city_of_gatineau"),
        "JOB1_BULLET_2": bullet_or_fallback(gatineau, 1, "city_of_gatineau"),
        "JOB1_BULLET_3": bullet_or_fallback(gatineau, 2, "city_of_gatineau"),
        "JOB2_HEADER": "Software Engineer | Precision OS\tJanuary 2022 \u2013 September 2023",
        "JOB2_BULLET_1": bullet_or_fallback(precision, 0, "precision_os"),
        "JOB2_BULLET_2": bullet_or_fallback(precision, 1, "precision_os"),
        "JOB2_BULLET_3": bullet_or_fallback(precision, 2, "precision_os"),
        "JOB3_HEADER": "Full-Stack Developer | Syntax (Consulting)\tJanuary 2021 \u2013 January 2022",
        "JOB3_BULLET_1": bullet_or_fallback(syntax, 0, "syntax"),
        "JOB3_BULLET_2": bullet_or_fallback(syntax, 1, "syntax"),
        "JOB3_BULLET_3": bullet_or_fallback(syntax, 2, "syntax"),
        "JOB4_HEADER": "Web Developer | University of Ottawa\tMay 2019 \u2013 January 2021",
        "JOB4_BULLET_1": bullet_or_fallback(uottawa, 0, "uottawa"),
        "JOB4_BULLET_2": bullet_or_fallback(uottawa, 1, "uottawa"),
        "JOB4_BULLET_3": bullet_or_fallback(uottawa, 2, "uottawa"),
        "EDUCATION": "Bachelor of Applied Science, Computer Engineering\t2016 \u2013 2020",
        "UNIVERSITY": "University of Ottawa",
    })
    
    skills_keys = skills_placeholders
    job_header_keys = ["JOB1_HEADER", "JOB2_HEADER", "JOB3_HEADER", "JOB4_HEADER"]
    
    _fill_template_placeholders(doc, replacements, skills_keys=skills_keys, job_header_keys=job_header_keys)
    
    # Post-process: turn "LinkedIn" text into a real hyperlink
    linkedin_url = profile.get("linkedin", "")
    if linkedin_url:
        _add_hyperlink_to_contact(doc, "LinkedIn", linkedin_url)
    
    company = job.get("company", "Unknown")
    title = job.get("title", "Unknown")
    job_url = job.get("url", "")
    safe_company = sanitize_filename(company)
    safe_title = sanitize_filename(_clean_job_title(title))
    
    # Create folder: CompanyName_JobTitle
    folder_name = f"{safe_company}_{safe_title}"
    job_folder = output_dir / folder_name
    job_folder.mkdir(exist_ok=True)
    
    # Save resume as AG_Resume.docx
    filepath = job_folder / "AG_Resume.docx"
    doc.save(str(filepath))
    
    # Save job link
    if job_url:
        link_path = job_folder / "job_link.txt"
        with open(link_path, "w", encoding="utf-8") as f:
            f.write(f"{title}\n{company}\n{job_url}\n")
    
    logger.info(f"Generated resume: {folder_name}/AG_Resume.docx")
    
    return filepath


def _generate_from_scratch(job: dict, tailored_data: dict, config: dict, output_dir: Path) -> Path | None:
    """Generate a resume from scratch (no template)."""
    
    profile = config["profile"]
    base_resume = config["base_resume"]
    
    company = job.get("company", "Unknown")
    title = job.get("title", "Unknown")
    
    doc = Document()
    
    # --- Styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # --- Header: Name ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(profile["name"].upper())
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "Calibri"
    name_para.space_after = Pt(2)
    
    # --- Header: Title ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Full-Stack Developer")
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(50, 50, 50)
    title_run.font.name = "Calibri"
    title_para.space_after = Pt(2)
    
    # --- Header: Contact ---
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact_para.add_run(
        f"{profile['email']}  |  LinkedIn  |  French (C2) & English (C2)"
    )
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(80, 80, 80)
    contact_run.font.name = "Calibri"
    contact_para.space_after = Pt(6)
    
    # --- Summary ---
    add_section_header(doc, "PROFESSIONAL SUMMARY")
    summary_text = tailored_data.get("summary", "")
    if not summary_text:
        summary_text = (
            f"Full-Stack Developer with {profile['years_experience']}+ years of experience "
            "building scalable web applications, REST APIs, and cloud-based solutions. "
            "Bilingual (French C2 / English C2) with expertise in JavaScript/TypeScript, "
            "React, Node.js, Python, and modern DevOps practices."
        )
    summary_para = doc.add_paragraph(summary_text)
    summary_para.style.font.size = Pt(10.5)
    summary_para.space_after = Pt(6)
    
    # --- Skills (dynamic) ---
    add_section_header(doc, "TECHNICAL SKILLS")
    
    skills = _get_tailored_skills(tailored_data)
    
    for label, skill_list in skills.items():
        p = doc.add_paragraph()
        p.space_after = Pt(1)
        p.space_before = Pt(1)
        label_run = p.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        label_run.font.name = "Calibri"
        skills_run = p.add_run(skill_list)
        skills_run.font.size = Pt(10)
        skills_run.font.name = "Calibri"
    
    # --- Experience ---
    add_section_header(doc, "PROFESSIONAL EXPERIENCE")
    
    roles = [
        {
            "title": "Full-Stack Developer",
            "company": "City of Gatineau",
            "dates": "September 2023 – Present",
            "key": "city_of_gatineau",
        },
        {
            "title": "Software Engineer",
            "company": "Precision OS",
            "dates": "January 2022 – September 2023",
            "key": "precision_os",
        },
        {
            "title": "Full-Stack Developer",
            "company": "Syntax (Consulting)",
            "dates": "January 2021 – January 2022",
            "key": "syntax",
        },
        {
            "title": "Web Developer",
            "company": "University of Ottawa",
            "dates": "May 2019 – January 2021",
            "key": "uottawa",
        },
    ]
    
    tailored_bullets = tailored_data.get("bullets", {})
    
    for role in roles:
        role_para = doc.add_paragraph()
        role_para.space_before = Pt(8)
        role_para.space_after = Pt(2)
        
        title_run = role_para.add_run(f"{role['title']} | {role['company']}")
        title_run.bold = True
        title_run.font.size = Pt(10.5)
        title_run.font.name = "Calibri"
        
        role_para.add_run("    ")
        
        date_run = role_para.add_run(role["dates"])
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(80, 80, 80)
        date_run.font.name = "Calibri"
        
        bullets = _get_bullets(tailored_bullets, role["key"], base_resume["bullets"].get(role["key"], []))
        
        for bullet_text in bullets:
            bp = doc.add_paragraph(style="List Bullet")
            bp.space_before = Pt(1)
            bp.space_after = Pt(1)
            br = bp.add_run(bullet_text)
            br.font.size = Pt(10)
            br.font.name = "Calibri"
    
    # --- Education ---
    add_section_header(doc, "EDUCATION")
    
    edu_para = doc.add_paragraph()
    edu_para.space_before = Pt(4)
    edu_title = edu_para.add_run("Bachelor of Applied Science, Computer Engineering")
    edu_title.bold = True
    edu_title.font.size = Pt(10.5)
    edu_title.font.name = "Calibri"
    edu_para.add_run("    ")
    edu_date = edu_para.add_run("2016 – 2020")
    edu_date.font.size = Pt(10)
    edu_date.font.color.rgb = RGBColor(80, 80, 80)
    edu_date.font.name = "Calibri"
    
    uni_para = doc.add_paragraph("University of Ottawa")
    uni_para.space_before = Pt(0)
    uni_run = uni_para.runs[0]
    uni_run.font.size = Pt(10)
    uni_run.font.color.rgb = RGBColor(80, 80, 80)
    uni_run.font.name = "Calibri"
    
    # --- Save ---
    safe_company = sanitize_filename(company)
    safe_title = sanitize_filename(_clean_job_title(title))
    job_url = job.get("url", "")
    
    folder_name = f"{safe_company}_{safe_title}"
    job_folder = output_dir / folder_name
    job_folder.mkdir(exist_ok=True)
    
    filepath = job_folder / "AG_Resume.docx"
    doc.save(str(filepath))
    
    if job_url:
        link_path = job_folder / "job_link.txt"
        with open(link_path, "w", encoding="utf-8") as f:
            f.write(f"{title}\n{company}\n{job_url}\n")
    
    logger.info(f"Generated resume: {folder_name}/AG_Resume.docx")
    
    return filepath


def add_section_header(doc, text: str):
    """Add a section header with bottom border."""
    para = doc.add_paragraph()
    para.space_before = Pt(12)
    para.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '000000',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_all_resumes(scored_jobs: list[dict], tailored: dict, config: dict, output_dir: Path) -> list[Path]:
    """Generate tailored .docx resumes for all tailored jobs."""
    if not HAS_DOCX:
        logger.error("python-docx not installed. Cannot generate resumes.")
        return []
    
    from ats_checker import validate_resume_file, format_report
    
    resumes_dir = output_dir / "resumes"
    resumes_dir.mkdir(exist_ok=True)
    
    generated = []
    ats_results = []
    
    for job in scored_jobs:
        job_url = job.get("url", "")
        if job_url not in tailored:
            continue
        
        filepath = generate_resume_docx(job, tailored[job_url], config, resumes_dir)
        if filepath:
            generated.append(filepath)
            
            job_desc = job.get("description", "")
            result = validate_resume_file(str(filepath), job_desc)
            overall = result.get("overall", {})
            score = overall.get("score", 0)
            ready = overall.get("ats_ready", False)
            
            status = "✅ ATS-READY" if ready else f"⚠️ SCORE:{score}%"
            logger.info(f"  ATS check [{status}]: {filepath.name}")
            
            if not ready and overall.get("critical_fails"):
                for issue in overall["critical_fails"]:
                    logger.warning(f"    🔴 {issue}")
            
            ats_results.append({"file": str(filepath), 
                                "job": job.get("title", ""), 
                                "job_url": job.get("url", ""), 
                                "result": result})
    
    # Save ATS report
    if ats_results:
        ats_report_lines = ["# ATS Validation Report\n"]
        for ar in ats_results:
            ats_report_lines.append(f"## {ar['job']}")
            ats_report_lines.append(f"Link: {ar['job_url']}")
            ats_report_lines.append(f"File: {ar['file']}")
            ats_report_lines.append(format_report(ar['result']))
        
        ats_report_path = output_dir / "ats_report.md"
        with open(ats_report_path, "w", encoding="utf-8") as f:
            f.write("\n".join(ats_report_lines))
        logger.info(f"ATS report saved: {ats_report_path}")
    
    logger.info(f"Generated {len(generated)} tailored resumes in {resumes_dir}")
    return generated